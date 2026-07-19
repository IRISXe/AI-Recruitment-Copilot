from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from unittest.mock import patch
from uuid import UUID, uuid4

import pytest
from docx import Document
from fastapi import UploadFile, status
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import Settings
from app.core.exceptions import AppException
from app.extraction.resume_extractor import EXTRACTOR_VERSION
from app.models.resume import Resume
from app.models.resume_content import ResumeContent
from app.services.resume_service import ResumeDownload


RESUMES_URL = "/api/v1/resumes"
CANDIDATES_URL = "/api/v1/candidates"

PDF_BYTES = b"%PDF-1.4\nHarsha Resume\n%%EOF"
DOCX_BYTES = b"PK\x03\x04test-docx-download-content"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

PDF_EXTRACTED_TEXT = (
    "Harsha Vardhan Backend Developer Python FastAPI PostgreSQL"
)



def valid_candidate_payload(
    *,
    full_name: str = "Harsha Vardhan",
    email: str = "harsha@example.com",
) -> dict[str, object]:
    return {
        "full_name": full_name,
        "email": email,
        "phone": "+91 9876543210",
        "current_location": "Hyderabad",
        "current_role": "Backend Developer",
        "total_experience_months": 18,
        "skills": ["Python", "FastAPI", "PostgreSQL"],
    }


def create_candidate_through_api(
    client: TestClient,
    *,
    full_name: str = "Harsha Vardhan",
    email: str = "harsha@example.com",
) -> dict[str, object]:
    response = client.post(
        CANDIDATES_URL,
        json=valid_candidate_payload(
            full_name=full_name,
            email=email,
        ),
    )

    assert response.status_code == 201

    return response.json()


def valid_resume_payload(
    *,
    candidate_id: str | None = None,
    stored_filename: str = "harsha-resume.pdf",
    is_primary: bool = False,
) -> dict[str, object]:
    return {
        "candidate_id": candidate_id or str(uuid4()),
        "original_filename": "Harsha_Resume.pdf",
        "stored_filename": stored_filename,
        "storage_path": (
            f"uploads/resumes/{stored_filename}"
        ),
        "content_type": "application/pdf",
        "file_size_bytes": 245760,
        "is_primary": is_primary,
    }


def create_resume_through_api(
    client: TestClient,
    *,
    candidate_id: str,
    stored_filename: str = "harsha-resume.pdf",
    is_primary: bool = False,
) -> dict[str, object]:
    response = client.post(
        RESUMES_URL,
        json=valid_resume_payload(
            candidate_id=candidate_id,
            stored_filename=stored_filename,
            is_primary=is_primary,
        ),
    )

    assert response.status_code == 201

    return response.json()


def build_uploaded_resume(
    *,
    candidate_id: UUID,
    is_primary: bool = False,
) -> Resume:
    timestamp = datetime.now(UTC)

    return Resume(
        id=uuid4(),
        candidate_id=candidate_id,
        original_filename="Harsha_Resume.pdf",
        stored_filename="stored-resume.pdf",
        storage_path=(
            "local_storage/resumes/stored-resume.pdf"
        ),
        content_type="application/pdf",
        file_size_bytes=len(PDF_BYTES),
        is_primary=is_primary,
        created_at=timestamp,
        updated_at=timestamp,
    )


def test_create_resume_returns_201_and_persists_resume(
    client: TestClient,
    db_session: Session,
) -> None:
    candidate = create_candidate_through_api(client)

    response = client.post(
        RESUMES_URL,
        json=valid_resume_payload(
            candidate_id=str(candidate["id"]),
        ),
    )

    assert response.status_code == 201

    body = response.json()
    resume_id = UUID(body["id"])

    assert body["candidate_id"] == candidate["id"]
    assert body["original_filename"] == "Harsha_Resume.pdf"
    assert body["stored_filename"] == "harsha-resume.pdf"
    assert (
        body["storage_path"]
        == "uploads/resumes/harsha-resume.pdf"
    )
    assert body["content_type"] == "application/pdf"
    assert body["file_size_bytes"] == 245760
    assert body["is_primary"] is False
    assert body["created_at"] is not None
    assert body["updated_at"] is not None

    db_session.expire_all()

    persisted_resume = db_session.get(
        Resume,
        resume_id,
    )

    assert persisted_resume is not None
    assert (
        persisted_resume.candidate_id
        == UUID(str(candidate["id"]))
    )
    assert (
        persisted_resume.original_filename
        == "Harsha_Resume.pdf"
    )
    assert (
        persisted_resume.stored_filename
        == "harsha-resume.pdf"
    )


@pytest.mark.parametrize(
    "invalid_payload",
    [
        {},
        {
            **valid_resume_payload(),
            "candidate_id": "not-a-valid-uuid",
        },
        {
            **valid_resume_payload(),
            "file_size_bytes": -1,
        },
    ],
)
def test_create_resume_rejects_invalid_body(
    client: TestClient,
    db_session: Session,
    invalid_payload: dict[str, object],
) -> None:
    response = client.post(
        RESUMES_URL,
        json=invalid_payload,
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )

    count = db_session.scalar(
        select(func.count()).select_from(Resume)
    )

    assert count == 0


def test_create_resume_returns_404_for_unknown_candidate(
    client: TestClient,
) -> None:
    response = client.post(
        RESUMES_URL,
        json=valid_resume_payload(
            candidate_id=str(uuid4()),
        ),
    )

    assert response.status_code == 404
    assert (
        response.json()["error"]["code"]
        == "candidate_not_found"
    )


def test_create_primary_resume_demotes_existing_primary(
    client: TestClient,
    db_session: Session,
) -> None:
    candidate = create_candidate_through_api(client)

    first = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="first-resume.pdf",
        is_primary=True,
    )

    second = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="second-resume.pdf",
        is_primary=True,
    )

    first_response = client.get(
        f"{RESUMES_URL}/{first['id']}"
    )
    second_response = client.get(
        f"{RESUMES_URL}/{second['id']}"
    )

    assert first_response.status_code == 200
    assert second_response.status_code == 200

    assert first_response.json()["is_primary"] is False
    assert second_response.json()["is_primary"] is True

    db_session.expire_all()

    first_resume = db_session.get(
        Resume,
        UUID(str(first["id"])),
    )
    second_resume = db_session.get(
        Resume,
        UUID(str(second["id"])),
    )

    assert first_resume is not None
    assert second_resume is not None
    assert first_resume.is_primary is False
    assert second_resume.is_primary is True


def test_list_resumes_returns_newest_with_pagination(
    client: TestClient,
    db_session: Session,
) -> None:
    candidate = create_candidate_through_api(client)

    first = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="first-resume.pdf",
    )
    second = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="second-resume.pdf",
    )
    third = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="third-resume.pdf",
    )

    db_session.expire_all()

    first_resume = db_session.get(
        Resume,
        UUID(str(first["id"])),
    )
    second_resume = db_session.get(
        Resume,
        UUID(str(second["id"])),
    )
    third_resume = db_session.get(
        Resume,
        UUID(str(third["id"])),
    )

    assert first_resume is not None
    assert second_resume is not None
    assert third_resume is not None

    first_resume.created_at = datetime(
        9999,
        1,
        1,
        tzinfo=UTC,
    )
    second_resume.created_at = datetime(
        9999,
        1,
        2,
        tzinfo=UTC,
    )
    third_resume.created_at = datetime(
        9999,
        1,
        3,
        tzinfo=UTC,
    )

    db_session.commit()

    first_page = client.get(
        RESUMES_URL,
        params={
            "offset": 0,
            "limit": 2,
        },
    )

    second_page = client.get(
        RESUMES_URL,
        params={
            "offset": 2,
            "limit": 1,
        },
    )

    assert first_page.status_code == 200
    assert second_page.status_code == 200

    assert [
        resume["id"]
        for resume in first_page.json()
    ] == [
        third["id"],
        second["id"],
    ]

    assert [
        resume["id"]
        for resume in second_page.json()
    ] == [
        first["id"],
    ]


@pytest.mark.parametrize(
    "query_string",
    [
        "?offset=-1&limit=20",
        "?offset=0&limit=0",
        "?offset=0&limit=101",
    ],
)
def test_list_resumes_rejects_invalid_pagination(
    client: TestClient,
    query_string: str,
) -> None:
    response = client.get(
        f"{RESUMES_URL}{query_string}"
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )
    assert (
        response.json()["error"]["details"][0]["loc"][0]
        == "query"
    )


def test_list_resumes_filters_by_candidate_id(
    client: TestClient,
) -> None:
    first_candidate = create_candidate_through_api(
        client,
        full_name="First Candidate",
        email="first@example.com",
    )
    second_candidate = create_candidate_through_api(
        client,
        full_name="Second Candidate",
        email="second@example.com",
    )

    first_resume = create_resume_through_api(
        client,
        candidate_id=str(first_candidate["id"]),
        stored_filename="first-candidate-resume.pdf",
    )

    create_resume_through_api(
        client,
        candidate_id=str(second_candidate["id"]),
        stored_filename="second-candidate-resume.pdf",
    )

    response = client.get(
        RESUMES_URL,
        params={
            "candidate_id": first_candidate["id"],
        },
    )

    assert response.status_code == 200
    assert len(response.json()) == 1

    assert response.json()[0]["id"] == first_resume["id"]
    assert (
        response.json()[0]["candidate_id"]
        == first_candidate["id"]
    )


def test_list_resumes_rejects_malformed_candidate_id(
    client: TestClient,
) -> None:
    response = client.get(
        RESUMES_URL,
        params={
            "candidate_id": "not-a-valid-uuid",
        },
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )
    assert response.json()["error"]["details"][0]["loc"] == [
        "query",
        "candidate_id",
    ]


def test_get_resume_by_id_returns_existing_resume(
    client: TestClient,
) -> None:
    candidate = create_candidate_through_api(client)

    resume = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
    )

    response = client.get(
        f"{RESUMES_URL}/{resume['id']}"
    )

    assert response.status_code == 200
    assert response.json()["id"] == resume["id"]
    assert (
        response.json()["candidate_id"]
        == candidate["id"]
    )
    assert (
        response.json()["stored_filename"]
        == "harsha-resume.pdf"
    )


def test_get_resume_by_id_returns_404_for_unknown_uuid(
    client: TestClient,
) -> None:
    response = client.get(
        f"{RESUMES_URL}/{uuid4()}"
    )

    assert response.status_code == 404
    assert (
        response.json()["error"]["code"]
        == "resume_not_found"
    )


def test_get_resume_by_id_rejects_malformed_uuid(
    client: TestClient,
) -> None:
    response = client.get(
        f"{RESUMES_URL}/not-a-valid-uuid"
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )
    assert response.json()["error"]["details"][0]["loc"] == [
        "path",
        "resume_id",
    ]


def test_download_resume_returns_pdf_file(
    client: TestClient,
    tmp_path: Path,
) -> None:
    resume_id = uuid4()

    file_path = tmp_path / "stored-resume.pdf"
    file_path.write_bytes(PDF_BYTES)

    download = ResumeDownload(
        file_path=file_path,
        filename="Harsha_Resume.pdf",
        content_type="application/pdf",
    )

    with patch(
        "app.api.routes.resumes.get_resume_download_service",
        return_value=download,
    ) as download_service:
        response = client.get(
            f"{RESUMES_URL}/{resume_id}/download"
        )

    download_service.assert_called_once()

    _, call_kwargs = download_service.call_args

    assert call_kwargs["resume_id"] == resume_id
    assert response.status_code == 200
    assert response.content == PDF_BYTES
    assert response.headers["content-type"] == "application/pdf"

    content_disposition = response.headers[
        "content-disposition"
    ]

    assert content_disposition.startswith("attachment;")
    assert (
        'filename="Harsha_Resume.pdf"'
        in content_disposition
    )


def test_download_resume_returns_docx_file(
    client: TestClient,
    tmp_path: Path,
) -> None:
    resume_id = uuid4()

    file_path = tmp_path / "stored-resume.docx"
    file_path.write_bytes(DOCX_BYTES)

    download = ResumeDownload(
        file_path=file_path,
        filename="Harsha_Resume.docx",
        content_type=DOCX_CONTENT_TYPE,
    )

    with patch(
        "app.api.routes.resumes.get_resume_download_service",
        return_value=download,
    ):
        response = client.get(
            f"{RESUMES_URL}/{resume_id}/download"
        )

    assert response.status_code == 200
    assert response.content == DOCX_BYTES
    assert (
        response.headers["content-type"]
        == DOCX_CONTENT_TYPE
    )

    content_disposition = response.headers[
        "content-disposition"
    ]

    assert content_disposition.startswith("attachment;")
    assert (
        'filename="Harsha_Resume.docx"'
        in content_disposition
    )


def test_download_resume_returns_404_for_unknown_uuid(
    client: TestClient,
) -> None:
    response = client.get(
        f"{RESUMES_URL}/{uuid4()}/download"
    )

    assert response.status_code == 404
    assert (
        response.json()["error"]["code"]
        == "resume_not_found"
    )


def test_download_resume_rejects_malformed_uuid(
    client: TestClient,
) -> None:
    with patch(
        "app.api.routes.resumes.get_resume_download_service",
    ) as download_service:
        response = client.get(
            f"{RESUMES_URL}/not-a-valid-uuid/download"
        )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )
    assert response.json()["error"]["details"][0]["loc"] == [
        "path",
        "resume_id",
    ]

    download_service.assert_not_called()


def test_download_resume_returns_404_for_missing_physical_file(
    client: TestClient,
) -> None:
    resume_id = uuid4()

    with patch(
        "app.api.routes.resumes.get_resume_download_service",
        side_effect=AppException(
            status_code=status.HTTP_404_NOT_FOUND,
            code="resume_file_not_found",
            message="The Resume file could not be found.",
        ),
    ):
        response = client.get(
            f"{RESUMES_URL}/{resume_id}/download"
        )

    assert response.status_code == 404
    assert response.json() == {
    "error": {
        "code": "resume_file_not_found",
        "message": "The Resume file could not be found.",
        "details": None,
    }
}


def test_download_resume_hides_unsafe_storage_path(
    client: TestClient,
) -> None:
    resume_id = uuid4()
    unsafe_path = (
        "C:/private/server-storage/resumes/secret.pdf"
    )

    with patch(
        "app.api.routes.resumes.get_resume_download_service",
        side_effect=AppException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            code="resume_download_failed",
            message=(
                "The Resume file could not be downloaded."
            ),
        ),
    ):
        response = client.get(
            f"{RESUMES_URL}/{resume_id}/download"
        )

    assert response.status_code == 500
    assert (
        response.json()["error"]["code"]
        == "resume_download_failed"
    )
    assert response.json()["error"]["message"] == (
        "The Resume file could not be downloaded."
    )
    assert unsafe_path not in response.text
    assert "server-storage" not in response.text


def test_update_resume_sets_primary_and_demotes_previous_primary(
    client: TestClient,
    db_session: Session,
) -> None:
    candidate = create_candidate_through_api(client)

    first = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="first-resume.pdf",
        is_primary=True,
    )

    second = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
        stored_filename="second-resume.pdf",
        is_primary=False,
    )

    response = client.patch(
        f"{RESUMES_URL}/{second['id']}",
        json={
            "is_primary": True,
        },
    )

    assert response.status_code == 200
    assert response.json()["is_primary"] is True
    assert response.json()["id"] == second["id"]

    first_response = client.get(
        f"{RESUMES_URL}/{first['id']}"
    )

    assert first_response.status_code == 200
    assert first_response.json()["is_primary"] is False

    db_session.expire_all()

    first_resume = db_session.get(
        Resume,
        UUID(str(first["id"])),
    )
    second_resume = db_session.get(
        Resume,
        UUID(str(second["id"])),
    )

    assert first_resume is not None
    assert second_resume is not None
    assert first_resume.is_primary is False
    assert second_resume.is_primary is True


@pytest.mark.parametrize(
    "invalid_payload",
    [
        {},
        {
            "is_primary": "not-a-boolean",
        },
    ],
)
def test_update_resume_rejects_invalid_body(
    client: TestClient,
    invalid_payload: dict[str, object],
) -> None:
    candidate = create_candidate_through_api(client)

    resume = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
    )

    response = client.patch(
        f"{RESUMES_URL}/{resume['id']}",
        json=invalid_payload,
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )


def test_update_resume_returns_404_for_unknown_uuid(
    client: TestClient,
) -> None:
    response = client.patch(
        f"{RESUMES_URL}/{uuid4()}",
        json={
            "is_primary": True,
        },
    )

    assert response.status_code == 404
    assert (
        response.json()["error"]["code"]
        == "resume_not_found"
    )


def test_update_resume_rejects_malformed_uuid(
    client: TestClient,
) -> None:
    response = client.patch(
        f"{RESUMES_URL}/not-a-valid-uuid",
        json={
            "is_primary": True,
        },
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )
    assert response.json()["error"]["details"][0]["loc"] == [
        "path",
        "resume_id",
    ]


def test_delete_resume_returns_204_and_removes_resume(
    client: TestClient,
    db_session: Session,
) -> None:
    candidate = create_candidate_through_api(client)

    resume = create_resume_through_api(
        client,
        candidate_id=str(candidate["id"]),
    )

    resume_id = UUID(str(resume["id"]))

    response = client.delete(
        f"{RESUMES_URL}/{resume_id}"
    )

    assert response.status_code == 204
    assert response.content == b""

    db_session.expire_all()

    assert db_session.get(Resume, resume_id) is None

    repeated_response = client.delete(
        f"{RESUMES_URL}/{resume_id}"
    )

    assert repeated_response.status_code == 404
    assert (
        repeated_response.json()["error"]["code"]
        == "resume_not_found"
    )


def test_delete_resume_rejects_malformed_uuid(
    client: TestClient,
) -> None:
    response = client.delete(
        f"{RESUMES_URL}/not-a-valid-uuid"
    )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )
    assert response.json()["error"]["details"][0]["loc"] == [
        "path",
        "resume_id",
    ]


def test_upload_resume_accepts_multipart_file_and_returns_201(
    client: TestClient,
) -> None:
    expected_candidate_id = uuid4()

    resume = build_uploaded_resume(
        candidate_id=expected_candidate_id,
        is_primary=True,
    )

    captured: dict[str, object] = {}

    def fake_upload_service(
        session: Session,
        *,
        candidate_id: UUID,
        uploaded_file: UploadFile,
        is_primary: bool,
    ) -> Resume:
        captured["session"] = session
        captured["candidate_id"] = candidate_id
        captured["filename"] = uploaded_file.filename
        captured["content_type"] = uploaded_file.content_type
        captured["content"] = uploaded_file.file.read()
        captured["is_primary"] = is_primary

        return resume

    with patch(
        "app.api.routes.resumes.upload_resume_service",
        side_effect=fake_upload_service,
    ) as upload_service:
        response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": str(expected_candidate_id),
                "is_primary": "true",
            },
            files={
                "file": (
                    "Harsha_Resume.pdf",
                    PDF_BYTES,
                    "application/pdf",
                ),
            },
        )

    assert response.status_code == 201

    upload_service.assert_called_once()

    assert captured["candidate_id"] == expected_candidate_id
    assert captured["filename"] == "Harsha_Resume.pdf"
    assert captured["content_type"] == "application/pdf"
    assert captured["content"] == PDF_BYTES
    assert captured["is_primary"] is True

    body = response.json()

    assert body["id"] == str(resume.id)
    assert body["candidate_id"] == str(expected_candidate_id)
    assert body["original_filename"] == "Harsha_Resume.pdf"
    assert body["stored_filename"] == "stored-resume.pdf"
    assert body["content_type"] == "application/pdf"
    assert body["file_size_bytes"] == len(PDF_BYTES)
    assert body["is_primary"] is True


def test_upload_resume_defaults_is_primary_to_false(
    client: TestClient,
) -> None:
    candidate_id = uuid4()

    resume = build_uploaded_resume(
        candidate_id=candidate_id,
    )

    with patch(
        "app.api.routes.resumes.upload_resume_service",
        return_value=resume,
    ) as upload_service:
        response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": str(candidate_id),
            },
            files={
                "file": (
                    "Harsha_Resume.pdf",
                    PDF_BYTES,
                    "application/pdf",
                ),
            },
        )

    assert response.status_code == 201

    upload_service.assert_called_once()

    _, call_kwargs = upload_service.call_args

    assert call_kwargs["candidate_id"] == candidate_id
    assert call_kwargs["is_primary"] is False
    assert (
        call_kwargs["uploaded_file"].filename
        == "Harsha_Resume.pdf"
    )

    assert response.json()["is_primary"] is False


def test_upload_resume_rejects_malformed_candidate_uuid(
    client: TestClient,
) -> None:
    with patch(
        "app.api.routes.resumes.upload_resume_service",
    ) as upload_service:
        response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": "not-a-valid-uuid",
                "is_primary": "false",
            },
            files={
                "file": (
                    "Harsha_Resume.pdf",
                    PDF_BYTES,
                    "application/pdf",
                ),
            },
        )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )

    details = response.json()["error"]["details"]

    assert any(
        detail["loc"] == ["body", "candidate_id"]
        for detail in details
    )

    upload_service.assert_not_called()


def test_upload_resume_requires_candidate_id(
    client: TestClient,
) -> None:
    with patch(
        "app.api.routes.resumes.upload_resume_service",
    ) as upload_service:
        response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "is_primary": "false",
            },
            files={
                "file": (
                    "Harsha_Resume.pdf",
                    PDF_BYTES,
                    "application/pdf",
                ),
            },
        )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )

    details = response.json()["error"]["details"]

    assert any(
        detail["loc"] == ["body", "candidate_id"]
        for detail in details
    )

    upload_service.assert_not_called()


def test_upload_resume_requires_file(
    client: TestClient,
) -> None:
    with patch(
        "app.api.routes.resumes.upload_resume_service",
    ) as upload_service:
        response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": str(uuid4()),
                "is_primary": "false",
            },
        )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )

    details = response.json()["error"]["details"]

    assert any(
        detail["loc"] == ["body", "file"]
        for detail in details
    )

    upload_service.assert_not_called()


def test_upload_download_and_delete_resume_manages_real_file(
    client: TestClient,
    db_session: Session,
    tmp_path: Path,
) -> None:
    candidate = create_candidate_through_api(client)

    settings = Settings(
        database_url=(
            "postgresql+psycopg://"
            "user:password@localhost/test"
        ),
        resume_storage_directory=(
            tmp_path / "resumes"
        ),
    )

    with patch(
        "app.services.resume_service.get_settings",
        return_value=settings,
    ):
        upload_response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": str(candidate["id"]),
                "is_primary": "true",
            },
            files={
                "file": (
                    "Harsha_Resume.pdf",
                    PDF_BYTES,
                    "application/pdf",
                ),
            },
        )

        assert upload_response.status_code == 201

        uploaded_resume = upload_response.json()
        resume_id = UUID(uploaded_resume["id"])
        stored_path = Path(
            uploaded_resume["storage_path"]
        )

        assert stored_path.exists()
        assert stored_path.is_file()
        assert stored_path.read_bytes() == PDF_BYTES

        assert (
            stored_path.parent
            == settings.resume_storage_directory
        )

        assert (
            uploaded_resume["candidate_id"]
            == candidate["id"]
        )
        assert (
            uploaded_resume["original_filename"]
            == "Harsha_Resume.pdf"
        )
        assert (
            uploaded_resume["content_type"]
            == "application/pdf"
        )
        assert (
            uploaded_resume["file_size_bytes"]
            == len(PDF_BYTES)
        )
        assert uploaded_resume["is_primary"] is True

        db_session.expire_all()

        persisted_resume = db_session.get(
            Resume,
            resume_id,
        )

        assert persisted_resume is not None
        assert (
            persisted_resume.storage_path
            == stored_path.as_posix()
        )

        download_response = client.get(
            f"{RESUMES_URL}/{resume_id}/download"
        )

        assert download_response.status_code == 200
        assert download_response.content == PDF_BYTES
        assert (
            download_response.headers["content-type"]
            == "application/pdf"
        )
        assert (
            'filename="Harsha_Resume.pdf"'
            in download_response.headers["content-disposition"]
        )

        delete_response = client.delete(
            f"{RESUMES_URL}/{resume_id}"
        )

        assert delete_response.status_code == 204
        assert delete_response.content == b""

        assert not stored_path.exists()

        db_session.expire_all()

        assert (
            db_session.get(
                Resume,
                resume_id,
            )
            is None
        )


def build_completed_resume_content(

    *,
    resume_id: UUID,
    extracted_text: str = "Backend Developer",
) -> ResumeContent:
    timestamp = datetime.now(UTC)

    return ResumeContent(
        id=uuid4(),
        resume_id=resume_id,
        extracted_text=extracted_text,
        extraction_status="completed",
        extraction_error=None,
        extractor_version=EXTRACTOR_VERSION,
        extracted_at=timestamp,
        created_at=timestamp,
        updated_at=timestamp,
    )


def build_extractable_pdf_bytes() -> bytes:
    buffer = BytesIO()
    writer = PdfWriter()

    page = writer.add_blank_page(
        width=612,
        height=792,
    )

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )

    font_reference = writer._add_object(font)

    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): font_reference,
                }
            ),
        }
    )

    page[NameObject("/Resources")] = resources

    content_stream = DecodedStreamObject()
    content_stream.set_data(
        (
            "BT "
            "/F1 12 Tf "
            "72 720 Td "
            f"({PDF_EXTRACTED_TEXT}) Tj "
            "ET"
        ).encode("latin-1")
    )

    page[NameObject("/Contents")] = writer._add_object(
        content_stream
    )

    writer.write(buffer)

    return buffer.getvalue()


def build_extractable_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()

    document.add_paragraph("Harsha Vardhan")
    document.add_paragraph("Frontend Developer")
    document.add_paragraph("React TypeScript FastAPI")

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.cell(0, 0).text = "Experience"
    table.cell(0, 1).text = "18 months"

    document.save(buffer)

    return buffer.getvalue()


def test_extract_resume_content_returns_completed_content(
    client: TestClient,
) -> None:
    resume_id = uuid4()

    content = build_completed_resume_content(
        resume_id=resume_id,
        extracted_text="Python FastAPI PostgreSQL",
    )

    with patch(
        "app.api.routes.resumes.extract_resume_content_service",
        return_value=content,
    ) as extract_service:
        response = client.post(
            f"{RESUMES_URL}/{resume_id}/extract"
        )

    assert response.status_code == 200

    extract_service.assert_called_once()

    _, call_kwargs = extract_service.call_args

    assert call_kwargs["resume_id"] == resume_id

    body = response.json()

    assert body["id"] == str(content.id)
    assert body["resume_id"] == str(resume_id)
    assert body["extracted_text"] == "Python FastAPI PostgreSQL"
    assert body["extraction_status"] == "completed"
    assert body["extraction_error"] is None
    assert body["extractor_version"] == EXTRACTOR_VERSION
    assert body["extracted_at"] is not None
    assert body["created_at"] is not None
    assert body["updated_at"] is not None


@pytest.mark.parametrize(
    (
        "expected_status",
        "expected_code",
        "expected_message",
    ),
    [
        (
            status.HTTP_404_NOT_FOUND,
            "resume_not_found",
            "The requested resume does not exist.",
        ),
        (
            status.HTTP_422_UNPROCESSABLE_CONTENT,
            "resume_extraction_failed",
            "The Resume text could not be extracted.",
        ),
    ],
)
def test_extract_resume_content_returns_service_error(
    client: TestClient,
    expected_status: int,
    expected_code: str,
    expected_message: str,
) -> None:
    resume_id = uuid4()

    with patch(
        "app.api.routes.resumes.extract_resume_content_service",
        side_effect=AppException(
            status_code=expected_status,
            code=expected_code,
            message=expected_message,
        ),
    ):
        response = client.post(
            f"{RESUMES_URL}/{resume_id}/extract"
        )

    assert response.status_code == expected_status
    assert response.json() == {
        "error": {
            "code": expected_code,
            "message": expected_message,
            "details": None,
        }
    }


def test_extract_resume_content_rejects_malformed_uuid(
    client: TestClient,
) -> None:
    with patch(
        "app.api.routes.resumes.extract_resume_content_service",
    ) as extract_service:
        response = client.post(
            f"{RESUMES_URL}/not-a-valid-uuid/extract"
        )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )

    assert response.json()["error"]["details"][0]["loc"] == [
        "path",
        "resume_id",
    ]

    extract_service.assert_not_called()


def test_get_resume_content_returns_completed_content(
    client: TestClient,
) -> None:
    resume_id = uuid4()

    content = build_completed_resume_content(
        resume_id=resume_id,
        extracted_text="React TypeScript FastAPI",
    )

    with patch(
        "app.api.routes.resumes.get_resume_content_service",
        return_value=content,
    ) as get_content_service:
        response = client.get(
            f"{RESUMES_URL}/{resume_id}/content"
        )

    assert response.status_code == 200

    get_content_service.assert_called_once()

    _, call_kwargs = get_content_service.call_args

    assert call_kwargs["resume_id"] == resume_id

    body = response.json()

    assert body["id"] == str(content.id)
    assert body["resume_id"] == str(resume_id)
    assert body["extracted_text"] == "React TypeScript FastAPI"
    assert body["extraction_status"] == "completed"
    assert body["extraction_error"] is None
    assert body["extractor_version"] == EXTRACTOR_VERSION


def test_get_resume_content_returns_404_when_not_extracted(
    client: TestClient,
) -> None:
    resume_id = uuid4()

    with patch(
        "app.api.routes.resumes.get_resume_content_service",
        side_effect=AppException(
            status_code=status.HTTP_404_NOT_FOUND,
            code="resume_content_not_found",
            message=(
                "Extracted content does not exist "
                "for the requested resume."
            ),
        ),
    ):
        response = client.get(
            f"{RESUMES_URL}/{resume_id}/content"
        )

    assert response.status_code == 404
    assert response.json() == {
        "error": {
            "code": "resume_content_not_found",
            "message": (
                "Extracted content does not exist "
                "for the requested resume."
            ),
            "details": None,
        }
    }


def test_get_resume_content_rejects_malformed_uuid(
    client: TestClient,
) -> None:
    with patch(
        "app.api.routes.resumes.get_resume_content_service",
    ) as get_content_service:
        response = client.get(
            f"{RESUMES_URL}/not-a-valid-uuid/content"
        )

    assert response.status_code == 422
    assert (
        response.json()["error"]["code"]
        == "validation_error"
    )

    assert response.json()["error"]["details"][0]["loc"] == [
        "path",
        "resume_id",
    ]

    get_content_service.assert_not_called()


def test_upload_extract_reextract_and_get_real_pdf_content(
    client: TestClient,
    db_session: Session,
    tmp_path: Path,
) -> None:
    candidate = create_candidate_through_api(client)
    pdf_bytes = build_extractable_pdf_bytes()

    settings = Settings(
        database_url=(
            "postgresql+psycopg://"
            "user:password@localhost/test"
        ),
        resume_storage_directory=(
            tmp_path / "pdf-resumes"
        ),
    )

    with (
        patch(
            "app.services.resume_service.get_settings",
            return_value=settings,
        ),
        patch(
            "app.services.resume_content_service.get_settings",
            return_value=settings,
        ),
    ):
        upload_response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": str(candidate["id"]),
                "is_primary": "true",
            },
            files={
                "file": (
                    "Harsha_Resume.pdf",
                    pdf_bytes,
                    "application/pdf",
                ),
            },
        )

        assert upload_response.status_code == 201

        resume_id = UUID(
            upload_response.json()["id"]
        )

        first_extract_response = client.post(
            f"{RESUMES_URL}/{resume_id}/extract"
        )

        second_extract_response = client.post(
            f"{RESUMES_URL}/{resume_id}/extract"
        )

        content_response = client.get(
            f"{RESUMES_URL}/{resume_id}/content"
        )

    assert first_extract_response.status_code == 200
    assert second_extract_response.status_code == 200
    assert content_response.status_code == 200

    first_body = first_extract_response.json()
    second_body = second_extract_response.json()
    content_body = content_response.json()

    assert first_body["resume_id"] == str(resume_id)
    assert first_body["extraction_status"] == "completed"
    assert first_body["extracted_text"] == PDF_EXTRACTED_TEXT
    assert first_body["extraction_error"] is None
    assert first_body["extractor_version"] == EXTRACTOR_VERSION

    assert second_body["id"] == first_body["id"]
    assert second_body["extraction_status"] == "completed"
    assert second_body["extracted_text"] == PDF_EXTRACTED_TEXT

    assert content_body["id"] == first_body["id"]
    assert content_body["resume_id"] == str(resume_id)
    assert content_body["extracted_text"] == PDF_EXTRACTED_TEXT
    assert content_body["extraction_status"] == "completed"

    db_session.expire_all()

    persisted_content = db_session.scalar(
        select(ResumeContent).where(
            ResumeContent.resume_id == resume_id
        )
    )

    assert persisted_content is not None
    assert persisted_content.id == UUID(first_body["id"])
    assert persisted_content.extracted_text == PDF_EXTRACTED_TEXT
    assert persisted_content.extraction_status == "completed"
    assert persisted_content.extraction_error is None


def test_upload_extract_and_get_real_docx_content(
    client: TestClient,
    db_session: Session,
    tmp_path: Path,
) -> None:
    candidate = create_candidate_through_api(client)
    docx_bytes = build_extractable_docx_bytes()

    settings = Settings(
        database_url=(
            "postgresql+psycopg://"
            "user:password@localhost/test"
        ),
        resume_storage_directory=(
            tmp_path / "docx-resumes"
        ),
    )

    with (
        patch(
            "app.services.resume_service.get_settings",
            return_value=settings,
        ),
        patch(
            "app.services.resume_content_service.get_settings",
            return_value=settings,
        ),
    ):
        upload_response = client.post(
            f"{RESUMES_URL}/upload",
            data={
                "candidate_id": str(candidate["id"]),
                "is_primary": "false",
            },
            files={
                "file": (
                    "Harsha_Resume.docx",
                    docx_bytes,
                    DOCX_CONTENT_TYPE,
                ),
            },
        )

        assert upload_response.status_code == 201

        resume_id = UUID(
            upload_response.json()["id"]
        )

        extract_response = client.post(
            f"{RESUMES_URL}/{resume_id}/extract"
        )

        content_response = client.get(
            f"{RESUMES_URL}/{resume_id}/content"
        )

    assert extract_response.status_code == 200
    assert content_response.status_code == 200

    extracted_body = extract_response.json()
    content_body = content_response.json()

    assert extracted_body["resume_id"] == str(resume_id)
    assert extracted_body["extraction_status"] == "completed"
    assert extracted_body["extraction_error"] is None
    assert extracted_body["extractor_version"] == EXTRACTOR_VERSION

    extracted_text = extracted_body["extracted_text"]

    assert "Harsha Vardhan" in extracted_text
    assert "Frontend Developer" in extracted_text
    assert "React TypeScript FastAPI" in extracted_text
    assert "Experience" in extracted_text
    assert "18 months" in extracted_text

    assert content_body["id"] == extracted_body["id"]
    assert content_body["extracted_text"] == extracted_text
    assert content_body["extraction_status"] == "completed"

    db_session.expire_all()

    persisted_content = db_session.scalar(
        select(ResumeContent).where(
            ResumeContent.resume_id == resume_id
        )
    )

    assert persisted_content is not None
    assert persisted_content.extracted_text == extracted_text
    assert persisted_content.extraction_status == "completed"
