from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from app.extraction.resume_extractor import (
    DOCX_CONTENT_TYPE,
    PDF_CONTENT_TYPE,
    EmptyResumeTextError,
    ResumeDocumentReadError,
    UnsupportedResumeContentTypeError,
    extract_resume_text,
)


def test_extract_resume_text_extracts_and_normalizes_pdf_text() -> None:
    file_path = Path("stored-resume.pdf")

    first_page = MagicMock()
    first_page.extract_text.return_value = (
        "  Harsha   Vardhan  \n\n"
        "Frontend    Developer  "
    )

    second_page = MagicMock()
    second_page.extract_text.return_value = (
        "React     TypeScript"
    )

    reader = MagicMock()
    reader.is_encrypted = False
    reader.pages = [
        first_page,
        second_page,
    ]

    with patch(
        "app.extraction.resume_extractor.PdfReader",
        return_value=reader,
    ) as pdf_reader:
        result = extract_resume_text(
            file_path=file_path,
            content_type=PDF_CONTENT_TYPE,
        )

    pdf_reader.assert_called_once_with(
        file_path,
    )

    first_page.extract_text.assert_called_once_with()
    second_page.extract_text.assert_called_once_with()

    assert result == (
        "Harsha Vardhan\n"
        "Frontend Developer\n"
        "React TypeScript"
    )


def test_extract_resume_text_rejects_encrypted_pdf() -> None:
    file_path = Path("encrypted-resume.pdf")

    reader = MagicMock()
    reader.is_encrypted = True

    with patch(
        "app.extraction.resume_extractor.PdfReader",
        return_value=reader,
    ):
        with pytest.raises(
            ResumeDocumentReadError,
            match="Encrypted PDF Resume files are not supported",
        ):
            extract_resume_text(
                file_path=file_path,
                content_type=PDF_CONTENT_TYPE,
            )


def test_extract_resume_text_maps_unreadable_pdf() -> None:
    file_path = Path("corrupted-resume.pdf")

    with patch(
        "app.extraction.resume_extractor.PdfReader",
        side_effect=ValueError(
            "invalid PDF structure"
        ),
    ):
        with pytest.raises(
            ResumeDocumentReadError,
            match="The PDF Resume could not be read",
        ):
            extract_resume_text(
                file_path=file_path,
                content_type=PDF_CONTENT_TYPE,
            )


def test_extract_resume_text_rejects_empty_pdf() -> None:
    file_path = Path("empty-resume.pdf")

    first_page = MagicMock()
    first_page.extract_text.return_value = None

    second_page = MagicMock()
    second_page.extract_text.return_value = "   \n\n   "

    reader = MagicMock()
    reader.is_encrypted = False
    reader.pages = [
        first_page,
        second_page,
    ]

    with patch(
        "app.extraction.resume_extractor.PdfReader",
        return_value=reader,
    ):
        with pytest.raises(
            EmptyResumeTextError,
            match="does not contain extractable text",
        ):
            extract_resume_text(
                file_path=file_path,
                content_type=PDF_CONTENT_TYPE,
            )


def test_extract_resume_text_extracts_docx_paragraphs_and_tables(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "stored-resume.docx"

    document = Document()

    document.add_paragraph(
        "  Harsha   Vardhan  "
    )

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "React"

    table.cell(1, 0).text = "Experience"
    table.cell(1, 1).text = "18 months"

    document.add_paragraph(
        "Frontend    Developer"
    )

    document.save(
        file_path,
    )

    result = extract_resume_text(
        file_path=file_path,
        content_type=DOCX_CONTENT_TYPE,
    )

    assert result == (
        "Harsha Vardhan\n"
        "Skill | React\n"
        "Experience | 18 months\n"
        "Frontend Developer"
    )


def test_extract_resume_text_rejects_empty_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "empty-resume.docx"

    document = Document()
    document.save(
        file_path,
    )

    with pytest.raises(
        EmptyResumeTextError,
        match="does not contain extractable text",
    ):
        extract_resume_text(
            file_path=file_path,
            content_type=DOCX_CONTENT_TYPE,
        )


def test_extract_resume_text_maps_corrupted_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "corrupted-resume.docx"

    file_path.write_bytes(
        b"this is not a valid docx document"
    )

    with pytest.raises(
        ResumeDocumentReadError,
        match="The DOCX Resume could not be read",
    ):
        extract_resume_text(
            file_path=file_path,
            content_type=DOCX_CONTENT_TYPE,
        )


def test_extract_resume_text_rejects_unsupported_content_type() -> None:
    with pytest.raises(
        UnsupportedResumeContentTypeError,
        match="content type is not supported",
    ):
        extract_resume_text(
            file_path=Path("stored-resume.txt"),
            content_type="text/plain",
        )